"""
analyze_newicks.py

Uso:
    python analyze_newicks.py /caminho/para/newicks_dir

Saídas:
    - clados_summary.csv
    - clados_summary.docx
    - plain_newicks.txt  (concatenação dos 22 newicks)
"""

import os
import sys
import glob
from collections import defaultdict
from typing import Tuple, Set, List

import pandas as pd

# Dependência para parse Newick:
# Biopython Phylo é usado aqui
from Bio import Phylo
from Bio.Phylo.Newick import Clade

# Para exportar .docx
from docx import Document


def normalize_label(l: str) -> str:
    """Remove aspas externas e espaços."""
    if l is None:
        return ""
    return l.strip().strip("'\"")


def min_distance_to_leaf(clade: Clade) -> int:
    """
    Retorna o menor número de arestas entre este nó e qualquer folha
    (0 para folha).
    """
    if clade.is_terminal():
        return 0
    # recursivo
    child_dists = [min_distance_to_leaf(c) for c in clade.clades if c is not None]
    if not child_dists:
        return 0
    return 1 + min(child_dists)


def gather_terminals(clade: Clade) -> Tuple[str, ...]:
    """Retorna tupla ordenada com os nomes das folhas sob esse clado."""
    names = [normalize_label(term.name) for term in clade.get_terminals()]
    # Filtra vazios (caso algum leaf sem nome)
    names = [n for n in names if n]
    return tuple(sorted(names))


def analyze_newicks(path: str, metadata_csv: str = None,
                    threshold_3 = 0.20, threshold_2 = 0.40,
                    save_prefix: str = "clados_summary"):
    """
    path: diretório contendo os *.newick
    metadata_csv: opcional, CSV com colunas [filename, region_sigla] para mapear nomes de arquivo para regiões
    thresholds: critérios (padrão conforme solicitado)
    """

    # 1) Ler arquivos .newick
    files = sorted(glob.glob(os.path.join(path, "*.newick")))
    if len(files) == 0:
        raise FileNotFoundError(f"Nenhum arquivo .newick encontrado em {path}")

    n_files = len(files)
    plain_newicks_list: List[str] = []
    filenames = [os.path.basename(f) for f in files]

    # opcional: metadata para mapear filename->region_sigla
    file_to_region = {}
    if metadata_csv and os.path.exists(metadata_csv):
        md = pd.read_csv(metadata_csv, dtype=str)
        # espera colunas: filename, region
        for _, r in md.iterrows():
            file_to_region[r['filename']] = r.get('region', r['filename'])
    # por padrão, region = filename (poderá substituir manualmente depois)
    for fn in filenames:
        file_to_region.setdefault(fn, fn)

    # estruturas para contagem
    clade_stats = defaultdict(lambda: {
        "presence_count": 0,   # arquivos onde aparece (em qualquer nível)
        "level2_count": 0,     # arquivos onde aparece em 2º nível
        "level3_count": 0,     # arquivos onde aparece em 3º nível
        "files": set()         # nomes de arquivos onde apareceu
    })

    # percorre arquivos
    for fp in files:
        fname = os.path.basename(fp)
        with open(fp, "r", encoding="utf-8") as f:
            newick = f.read().strip()
        plain_newicks_list.append(newick)

        # Tenta parse com Biopython
        try:
            tree = Phylo.read(fp, "newick")
        except Exception as e:
            # fallback: tentar parse da string direta
            from io import StringIO
            tree = Phylo.read(StringIO(newick), "newick")

        # sets por arquivo (para não contar múltiplas vezes o mesmo clado no mesmo arquivo)
        seen_clades_in_file: Set[Tuple[str, ...]] = set()
        seen_level2_in_file: Set[Tuple[str, ...]] = set()
        seen_level3_in_file: Set[Tuple[str, ...]] = set()

        # percorre todos os nós
        for node in tree.find_clades(order='preorder'):
            # ignore pure terminal single-leaf clades
            terminals = gather_terminals(node)
            if len(terminals) < 2:
                continue
            # calcula nível (a partir das folhas)
            min_dist = min_distance_to_leaf(node)
            level = min_dist + 1  # folha -> min_dist=0 -> level=1

            # marca presença
            seen_clades_in_file.add(terminals)
            if level == 2:
                seen_level2_in_file.add(terminals)
            if level == 3:
                seen_level3_in_file.add(terminals)

        # atualiza contadores globais com sets locais
        for cl in seen_clades_in_file:
            s = clade_stats[cl]
            s["presence_count"] += 1
            s["files"].add(fname)
        for cl in seen_level2_in_file:
            clade_stats[cl]["level2_count"] += 1
        for cl in seen_level3_in_file:
            clade_stats[cl]["level3_count"] += 1

    # compor plain_newicks (string)
    plain_newicks = "\n".join(plain_newicks_list)
    # salvar plain_newicks opcional
    with open(f"{save_prefix}_plain_newicks.txt", "w", encoding="utf-8") as fh:
        fh.write(plain_newicks)

    # 2) calcular frequências e montar tabela
    rows = []
    for cl, stats in clade_stats.items():
        presence = stats["presence_count"]
        level2 = stats["level2_count"]
        level3 = stats["level3_count"]
        files_list = sorted(list(stats["files"]))

        freq_2 = level2 / n_files
        freq_3 = level3 / n_files
        presence_frac = presence / n_files

        # critérios para Lista Principal
        in_lista_principal = (freq_3 >= threshold_3) and (freq_2 >= threshold_2)

        # Estabilidade qualitativa (usando max(freq_2,freq_3))
        maxf = max(freq_2, freq_3)
        if maxf >= 0.85:
            stability = "Muito Alta"
        elif maxf >= 0.7:
            stability = "Alta"
        elif maxf >= 0.5:
            stability = "Moderada"
        else:
            stability = "Baixa"

        # Observações de frequência
        obs = f"3ºN={freq_3:.3f};2ºN={freq_2:.3f};presença={presence}/{n_files}"

        # região/arquivos onde aparece (aqui usamos os nomes de arquivo; se tiver metadata, substitua)
        regions_analyzed = ",".join([file_to_region.get(fn, fn) for fn in files_list])

        rows.append({
            "Clado": ",".join(cl),
            "Ocorrências (Estratificações)": f"{presence} / {n_files}",
            "Regiões analisadas": regions_analyzed,
            "Observações de Frequência": obs,
            "Categoria": "—",               # campo reservado: preencher manualmente se tiver mapeamento variável->categoria
            "Clados Representativos": "Sim" if in_lista_principal else "Não",
            "Estabilidade Global": stability,
            "freq_2": freq_2,
            "freq_3": freq_3
        })

    # montar DataFrame ordenado por presença decrescente
    df = pd.DataFrame(rows)
    df = df.sort_values(by=["Ocorrências (Estratificações)"], ascending=False, key=lambda s: s.map(lambda x: int(x.split(" / ")[0]) if isinstance(x, str) else 0))

    # Salvar CSV
    csv_out = f"{save_prefix}.csv"
    df.to_csv(csv_out, index=False, encoding="utf-8")

    # Salvar DOCX (tabela)
    doc = Document()
    doc.add_heading("Tabela 1 — Síntese Geral de Clados Identificados", level=1)
    table = doc.add_table(rows=1, cols=7)
    hdr_cells = table.rows[0].cells
    headers = ["Clado", "Ocorrências (Estratificações)", "Regiões analisadas", "Observações de Frequência", "Categoria", "Clados Representativos", "Estabilidade Global"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h

    for _, r in df.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(r["Clado"])
        row_cells[1].text = str(r["Ocorrências (Estratificações)"])
        row_cells[2].text = str(r["Regiões analisadas"])
        row_cells[3].text = str(r["Observações de Frequência"])
        row_cells[4].text = str(r["Categoria"])
        row_cells[5].text = str(r["Clados Representativos"])
        row_cells[6].text = str(r["Estabilidade Global"])

    docx_out = f"{save_prefix}.docx"
    doc.save(docx_out)

    # Lista Principal (clados que satisfazem os critérios)
    lista_principal = df[(df["Clados Representativos"] == "Sim")].copy()

    print("Resumo:")
    print(f"  Arquivos lidos: {n_files}")
    print(f"  Clados detectados: {len(clade_stats)}")
    print(f"  CSV salvo em: {csv_out}")
    print(f"  DOCX salvo em: {docx_out}")
    print(f"  plain_newicks salvo em: {save_prefix}_plain_newicks.txt")
    print(f"  Clados na Lista Principal: {len(lista_principal)} (salve {csv_out} para ver detalhes)")

    return {
        "df": df,
        "lista_principal": lista_principal,
        "plain_newicks": plain_newicks,
        "csv_out": csv_out,
        "docx_out": docx_out
    }


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python analyze_newicks.py /caminho/para/dir_newicks [metadata.csv]")
        sys.exit(1)
    path = sys.argv[1]
    metadata = sys.argv[2] if len(sys.argv) > 2 else None
    analyze_newicks(path, metadata_csv=metadata)
